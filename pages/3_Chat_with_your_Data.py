"""
Chat with your Data

Upload PDF documents and ask questions about them. The AI will search through
your documents and provide answers based on the content.
"""


# =========================================================
# IMPORTS (Libraries we need)
# =========================================================

from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from pypdf import PdfReader
from PIL import Image
import docx
import io, base64
import pandas as pd

# Streamlit: Framework for building web apps with Python
import streamlit as st

# os: For file operations and environment variables
import os

# ChatOpenAI: Connects to OpenAI's GPT models (like ChatGPT)
# OpenAIEmbeddings: Converts text to vectors (numbers) for similarity search
from langchain_openai import ChatOpenAI, OpenAIEmbeddings

# ChatPromptTemplate: Template for formatting messages to the AI Custom prompts
from langchain_core.prompts import ChatPromptTemplate

# PyPDFLoader: Loads and reads PDF files
from langchain_community.document_loaders import PyPDFLoader
import tempfile

# FAISS: Fast similarity search database (stores document chunks as vectors)
from langchain_community.vectorstores import FAISS

# RecursiveCharacterTextSplitter: Splits long documents into smaller chunks
from langchain.text_splitter import RecursiveCharacterTextSplitter

# LangGraph: For building agentic workflows
from langgraph.graph import StateGraph, START, END
from typing_extensions import TypedDict
from typing import Literal

# =========================================================
# PAGE SETUP
# =========================================================

st.set_page_config(
    page_title="Chat with Documents",
    page_icon="📚",
    layout="wide"  # Use full width of browser
)

st.title("📚 Chat with your Data (Agentic RAG)")
st.caption("AI agent that intelligently searches and answers from your documents")


# =========================================================
# SESSION STATE
# =========================================================

if "openai_key" not in st.session_state:
    st.session_state.openai_key = ""  # Store OpenAI API key

if "vector_store" not in st.session_state:
    st.session_state.vector_store = None  # Store document database

if "llm" not in st.session_state:
    st.session_state.llm = None  # Store language model instance

if "rag_messages" not in st.session_state:
    st.session_state.rag_messages = []  # Store chat history

if "processed_files" not in st.session_state:
    st.session_state.processed_files = []  # Track which files we've processed
if "rag_agent" not in st.session_state:
    st.session_state.rag_agent = None  # Store the agentic RAG workflow

# =========================================================
# SIDEBAR
# =========================================================

with st.sidebar:
    st.subheader("🔑 API Keys")
    
    if st.session_state.openai_key:
        st.success("✅ OpenAI Connected")
        
        # Show agent capabilities if agent is created
        if st.session_state.rag_agent:
            st.subheader("🤖 Agent Capabilities")
            st.write("✅ **Search Documents**")
            st.write("✅ **Grade Relevance**")
            st.write("✅ **Rewrite Questions**")
            st.write("✅ **Generate Answers**")
        
        if st.button("Change API Keys"):
            # Reset everything to start fresh
            st.session_state.openai_key = ""
            st.session_state.vector_store = None
            st.session_state.rag_messages = []
            st.session_state.rag_agent = None  # Also reset agent
            st.rerun()


# =========================================================
# API KEY INPUT
# =========================================================

if not st.session_state.openai_key:
    # Show input form for API key
    api_key = st.text_input(
        "OpenAI API Key",
        type="password",  # Hide the key as user types
        placeholder="sk-proj-..."
    )
    
    if st.button("Connect"):
        # Validate key format
        if api_key and api_key.startswith("sk-"):
            st.session_state.openai_key = api_key
            st.rerun()
        else:
            st.error("❌ Invalid API key format")
    
    st.stop()  # Don't show rest of app until connected

def pdf_to_text(file) -> str:
    reader = PdfReader(io.BytesIO(file.getvalue()))
    return "\n".join((page.extract_text() or "") for page in reader.pages)

def docx_to_text(file) -> str:
    d = docx.Document(io.BytesIO(file.getvalue()))
    return "\n".join(p.text for p in d.paragraphs if p.text)

def txt_to_text(file) -> str:
    return file.getvalue().decode("utf-8", errors="ignore")

def image_to_vision_text(file) -> str:
    """Vision analysis: describe image + extract text + structure."""
    img_bytes = file.getvalue()
    b64 = base64.b64encode(img_bytes).decode("utf-8")
    data_url = f"data:{file.type};base64,{b64}"

    msg = HumanMessage(content=[
        {"type": "text", "text": (
            "Analyze this image thoroughly.\n"
            "1) Describe the scene in detail.\n"
            "2) Extract ALL visible text as accurately as possible.\n"
            "3) List key entities (people, objects, brands, places, dates).\n"
            "4) If it's a document/screenshot, reconstruct headings/bullets/tables.\n"
            "Return a clean structured markdown summary."
        )},
        {"type": "image_url", "image_url": {"url": data_url}},
    ])

    vision_llm = ChatOpenAI(
        model="gpt-4o-mini",
        temperature=0,
        api_key=st.session_state.openai_key
    )

    return vision_llm.invoke([msg]).content

def pdf_to_documents(file):
    reader = PdfReader(io.BytesIO(file.getvalue()))
    docs = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": file.name, "page": i, "type": "pdf"}
            ))
    return docs

def pdf_to_documents_fallback(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file.getvalue())
        tmp_path = tmp.name

    loader = PyPDFLoader(tmp_path)
    docs = loader.load()

    os.remove(tmp_path)
    return docs

def csv_to_text(file) -> str:
    df = pd.read_csv(file)
    return df.to_markdown(index=False)

def excel_to_text(file) -> str:
    xls = pd.ExcelFile(io.BytesIO(file.getvalue()))
    parts = []
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        parts.append(f"# Sheet: {sheet}\n" + df.to_markdown(index=False))
    return "\n\n".join(parts)

# =========================================================
# PDF UPLOAD AND PROCESSING
# =========================================================

uploaded_files = st.file_uploader(
    "Upload files (PDF, DOCX, TXT, Images)",
    type=["pdf", "docx", "txt", "csv", "xlsx", "xls", "png", "jpg", "jpeg", "webp"],
    accept_multiple_files=True
)

if uploaded_files:
    current_files = sorted([f"{f.name}:{f.size}" for f in uploaded_files])

    if st.session_state.processed_files != current_files:
        with st.spinner("Processing files..."):
            documents = []

            for file in uploaded_files:
                ext = file.name.split(".")[-1].lower()

                if ext == "pdf":
                    documents.extend(pdf_to_documents(file))

                elif ext == "docx":
                    text = docx_to_text(file)
                    documents.append(Document(page_content=text, metadata={"source": file.name, "type": "docx"}))

                elif ext == "txt":
                    text = txt_to_text(file)
                    documents.append(Document(page_content=text, metadata={"source": file.name, "type": "txt"}))

                elif ext == "csv":
                    text = csv_to_text(file)
                    documents.append(Document(page_content=text, metadata={"source": file.name, "type": "csv"}))

                elif ext in ["xlsx", "xls"]:
                    text = excel_to_text(file)
                    documents.append(Document(page_content=text, metadata={"source": file.name, "type": "excel"}))
                    
                elif ext in ["png", "jpg", "jpeg", "webp"]:
                    st.image(Image.open(file), caption=file.name, use_container_width=True)
                    with st.spinner(f"Analyzing image: {file.name} ..."):
                        vision_text = image_to_vision_text(file)
                
                    documents.append(Document(
                        page_content=vision_text,
                        metadata={"source": file.name, "type": "image"}
                    ))

            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
            chunks = text_splitter.split_documents(documents)

            embeddings = OpenAIEmbeddings(api_key=st.session_state.openai_key)
            st.session_state.vector_store = FAISS.from_documents(chunks, embeddings)

            st.session_state.llm = ChatOpenAI(
                model="gpt-4o-mini",
                temperature=0,
                api_key=st.session_state.openai_key
            )

            st.session_state.rag_messages = []
            st.session_state.processed_files = current_files
            st.session_state.rag_agent = None  # rebuild graph for new docs

            st.success(f"✅ Processed {len(uploaded_files)} file(s)!")

# =============================================================================
# CREATE AGENTIC RAG WORKFLOW
# =============================================================================

#if st.session_state.vector_store and not st.session_state.rag_agent:
if st.session_state.vector_store and st.session_state.rag_agent is None:

    # Define state structure for the workflow
    class AgentState(TypedDict):
        question: str
        documents: list
        generation: str
        steps: list
        rewrite_count: int

    # Node 1: Retrieve documents
    def retrieve_documents(state: AgentState):
        question = state["question"]
        retriever = st.session_state.vector_store.as_retriever()
        docs = retriever.invoke(question)
        return {
            "question": question,
            "documents": docs,
            "generation": state.get("generation", ""),
            "steps": state.get("steps", []) + ["📚 Retrieved documents"],
            "rewrite_count": state.get("rewrite_count", 0),
        }

    # Node 2: Grade relevance
    def grade_documents(state: AgentState) -> Literal["generate", "rewrite"]:
        docs = state["documents"]
        rewrite_count = state.get("rewrite_count", 0)

        if rewrite_count >= 2:
            return "generate"
        if not docs:
            return "rewrite"

        prompt = f"""Are these documents relevant to the question: "{state['question']}"?

Documents:
{docs[0].page_content[:500]}

Answer with just 'yes' or 'no'."""
        response = st.session_state.llm.invoke(prompt)
        return "generate" if "yes" in response.content.lower() else "rewrite"

    # Node 3: Rewrite question
    def rewrite_question(state: AgentState):
        rewrite_count = state.get("rewrite_count", 0) + 1
        question = state["question"]

        new_question = st.session_state.llm.invoke(
            f"Rewrite this question to be more specific and searchable:\n\n{question}"
        ).content

        return {
            "question": new_question,
            "documents": [],
            "generation": "",
            "steps": state["steps"] + [f"🔄 Rewrote question (attempt {rewrite_count}): {new_question}"],
            "rewrite_count": rewrite_count,
        }

    # Node 4: Generate answer
    def generate_answer(state: AgentState):
        question = state["question"]
        docs = state["documents"]

        if not docs:
            return {
                "question": question,
                "documents": docs,
                "generation": "I couldn't find relevant information in the documents.",
                "steps": state["steps"] + ["❌ No relevant documents found"],
                "rewrite_count": state.get("rewrite_count", 0),
            }

        context = "\n\n---\n\n".join((d.page_content or "") for d in docs[:5])

        prompt = ChatPromptTemplate.from_messages([
            ("system", "Answer the question using ONLY the provided context. Be concise and accurate."),
            ("human", "Question: {question}\n\nContext: {context}\n\nAnswer:")
        ])

        response = st.session_state.llm.invoke(
            prompt.format_messages(question=question, context=context)
        )

        return {
            "question": question,
            "documents": docs,
            "generation": response.content,
            "steps": state["steps"] + ["💬 Generated answer"],
            "rewrite_count": state.get("rewrite_count", 0),
        }

    workflow = StateGraph(AgentState)
    workflow.add_node("retrieve", retrieve_documents)
    workflow.add_node("grade", grade_documents)
    workflow.add_node("rewrite", rewrite_question)
    workflow.add_node("generate", generate_answer)

    workflow.add_edge(START, "retrieve")
    workflow.add_conditional_edges("retrieve", grade_documents)
    workflow.add_edge("rewrite", "retrieve")
    workflow.add_edge("generate", END)

    st.session_state.rag_agent = workflow.compile()


if st.session_state.vector_store:
    
    # Display chat history
    for message in st.session_state.rag_messages:
        with st.chat_message(message["role"]):
            st.write(message["content"])
    
    # Handle user input
    user_input = st.chat_input("Ask a question about your documents...")
    
    if user_input:
        # Save user message
        st.session_state.rag_messages.append({
            "role": "user",
            "content": user_input
        })
        
        with st.chat_message("user"):
            st.write(user_input)
        
        # Generate response
                # Generate response using agentic workflow
        with st.chat_message("assistant"):
            with st.spinner("Agent is working..."):
                
                # Run the agentic workflow
                result = st.session_state.rag_agent.invoke({
                   "question": user_input,
                   "documents": [],
                   "generation": "",
                    "steps": [],
                   "rewrite_count": 0
                })
                 
                # Show agent's reasoning process
                with st.expander("🤖 View Agent Process", expanded=False):
                    st.markdown("### What the agent did:")
                    for step in result["steps"]:
                        st.markdown(f"- {step}")
                
                # Display final answer
                st.write(result["generation"])
                
                # Save to history
                st.session_state.rag_messages.append({
                    "role": "assistant",
                    "content": result["generation"]
                })

else:
    st.info("📄 Please upload PDF documents to start chatting.")
